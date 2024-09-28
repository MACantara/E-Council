import os
import json
import fitz  # PyMuPDF
from flask import Flask, request, render_template, send_from_directory
from docx import Document
from werkzeug.utils import secure_filename
from datetime import datetime
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure the API key for the generative model
genai.configure(api_key=os.environ['GOOGLE_GEMINI_AI_API_KEY'])

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

# Create the uploads directory if it doesn't exist
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        form_data = request.form.to_dict()
        subject_title = form_data.get('subject_title', '')

        # Handle file upload
        file = request.files.get('additional_context')
        additional_context = ""
        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            if filename.endswith('.pdf'):
                additional_context = extract_text_from_pdf(file_path)
            else:
                with open(file_path, 'r') as f:
                    additional_context = f.read()

        # Combine subject_title and additional_context for context
        context = f"{subject_title}\n{additional_context}"

        # Generate text using Gemini AI API with combined context
        form_data['date'] = generate_text("Suggest a date and only output that date for the event in the format 'Month Day, Year' (e.g., January 20, 2024).", context)
        form_data['introductory_paragraph'] = generate_text("Generate an introductory paragraph for the event.", context)
        form_data['time'] = generate_text("Suggest a time and only output that time for the event in the format H:mm AM/PM without leading zeros and without seconds.", context)
        form_data['location'] = generate_text("Suggest and only output a single location for the event.", context)
        form_data['participants'] = generate_text("Generate/predict and only output the number of participants (inte) for the event.", context)
        form_data['budget'] = generate_text("Suggest and only output an overall budget for the event in PHP currency.", context)

        # Convert non-string values to strings
        form_data['participants'] = str(form_data['participants'])
        form_data['budget'] = str(form_data['budget'])

        response_filename = 'concept_paper.docx'
        response_filepath = os.path.join(app.config['UPLOAD_FOLDER'], response_filename)
        create_docx(form_data, response_filepath)
        return render_template('index.html', download_link=response_filename)
    return render_template('index.html')

def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def generate_text(prompt, context):
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(f"{prompt} Context: {context}")
    try:
        response_json = json.loads(response.text)
        return response_json
    except json.JSONDecodeError:
        return response.text

def create_docx(data, output_path):
    template_path = os.path.join('ms-word-templates', 'concept-paper-with-header-and-footer-template.docx')
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    replace_placeholder(doc, '{{ signatory_name_1 }}', data['signatory_name_1'])
    replace_placeholder(doc, '{{ signatory_position_1 }}', data['signatory_position_1'])
    replace_placeholder(doc, '{{ signatory_name_2 }}', data['signatory_name_2'])
    replace_placeholder(doc, '{{ signatory_position_2 }}', data['signatory_position_2'])
    replace_placeholder(doc, '{{ signatory_department_2 }}', data['signatory_department_2'])
    replace_placeholder(doc, '{{ signatory_name_3 }}', data['signatory_name_3'])
    replace_placeholder(doc, '{{ signatory_position_3 }}', data['signatory_position_3'])
    replace_placeholder(doc, '{{ signatory_department_3 }}', data['signatory_department_3'])
    replace_placeholder(doc, '{{ subject_title }}', data['subject_title'])
    replace_placeholder(doc, '{{ date }}', data['date'])
    replace_placeholder(doc, '{{ introductory_paragraph }}', data['introductory_paragraph'])
    replace_placeholder(doc, '{{ time }}', data['time'])
    replace_placeholder(doc, '{{ location }}', data['location'])
    replace_placeholder(doc, '{{ participants }}', data['participants'])
    replace_placeholder(doc, '{{ budget }}', data['budget'])
    replace_placeholder(doc, '{{ descriptions }}', data['descriptions'])
    replace_placeholder(doc, '{{ objective_of_the_activity_1 }}', data['objective_of_the_activity_1'])
    replace_placeholder(doc, '{{ objective_of_the_activity_2 }}', data['objective_of_the_activity_2'])
    replace_placeholder(doc, '{{ objective_of_the_activity_3 }}', data['objective_of_the_activity_3'])
    replace_placeholder(doc, '{{ objective_of_the_activity_4 }}', data['objective_of_the_activity_4'])
    replace_placeholder(doc, '{{ learning_outcome_1 }}', data['learning_outcome_1'])
    replace_placeholder(doc, '{{ learning_outcome_2 }}', data['learning_outcome_2'])
    replace_placeholder(doc, '{{ learning_outcome_3 }}', data['learning_outcome_3'])
    replace_placeholder(doc, '{{ learning_outcome_4 }}', data['learning_outcome_4'])
    replace_placeholder(doc, '{{ number_of_expected_students }}', data['number_of_expected_students'])
    replace_placeholder(doc, '{{ number_of_expected_faculty }}', data['number_of_expected_faculty'])
    replace_placeholder(doc, '{{ prepared_by_name }}', data['prepared_by_name'])
    replace_placeholder(doc, '{{ prepared_by_position }}', data['prepared_by_position'])
    replace_placeholder(doc, '{{ prepared_by_student_council }}', data['prepared_by_student_council'])
    replace_placeholder(doc, '{{ signed_and_reviewed_by_name }}', data['signed_and_reviewed_by_name'])
    replace_placeholder(doc, '{{ signed_and_reviewed_by_position }}', data['signed_and_reviewed_by_position'])
    replace_placeholder(doc, '{{ signed_and_reviewed_by_student_council }}', data['signed_and_reviewed_by_student_council'])

    doc.save(output_path)

def replace_placeholder(doc, placeholder, replacement):
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            replace_in_paragraph(paragraph, placeholder, replacement)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, placeholder, replacement)

def replace_in_paragraph(paragraph, placeholder, replacement):
    # Concatenate all runs' text
    full_text = ''.join(run.text for run in paragraph.runs)
    
    # Replace the placeholder
    if placeholder in full_text:
        full_text = full_text.replace(placeholder, replacement)
        
        # Clear the paragraph's runs
        for run in paragraph.runs:
            run.text = ''
        
        # Add the new text back into the paragraph
        paragraph.add_run(full_text)

@app.route('/uploads/<filename>')
def uploads(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    app.run(host="0.0.0.0", debug=True)